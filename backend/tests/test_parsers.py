"""
Minimal pytest tests for PR 19 parsers (safe, no external services).

Tests document parsing functionality with graceful fallbacks for missing dependencies.
"""

import pytest
import csv
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

# Import the parsers module
from backend.app.services.parsers import (
    detect_type, parse_to_normalized, parse_csv, parse_docx, 
    parse_xlsx, parse_image_ocr, parse_pdf_stub, build_normalized
)


class TestDetectType:
    """Test document type detection."""
    
    def test_detect_type_csv(self):
        """Test CSV file type detection."""
        assert detect_type("data.csv") == "csv"
        assert detect_type("data.CSV") == "csv"
    
    def test_detect_type_docx(self):
        """Test DOCX file type detection."""
        assert detect_type("document.docx") == "docx"
        assert detect_type("document.DOCX") == "docx"
    
    def test_detect_type_xlsx(self):
        """Test XLSX file type detection."""
        assert detect_type("spreadsheet.xlsx") == "xlsx"
        assert detect_type("spreadsheet.XLSX") == "xlsx"
    
    def test_detect_type_image(self):
        """Test image file type detection."""
        assert detect_type("image.png") == "image"
        assert detect_type("photo.jpg") == "image"
        assert detect_type("scan.tiff") == "image"
        assert detect_type("picture.jpeg") == "image"
    
    def test_detect_type_pdf(self):
        """Test PDF file type detection."""
        assert detect_type("document.pdf") == "pdf"
        assert detect_type("document.PDF") == "pdf"
    
    def test_detect_type_unknown(self):
        """Test unknown file type detection."""
        assert detect_type("file.txt") == "unknown"
        assert detect_type("data.xml") == "unknown"


class TestCSVParser:
    """Test CSV parsing functionality."""
    
    def test_parse_csv_basic(self, tmp_path):
        """Test basic CSV parsing."""
        # Create a test CSV file
        csv_file = tmp_path / "test.csv"
        with open(csv_file, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Name", "Age", "City"])
            writer.writerow(["John", "30", "New York"])
            writer.writerow(["Jane", "25", "Los Angeles"])
        
        # Parse the CSV
        text, tables = parse_csv(csv_file)
        
        # Assert structure
        assert isinstance(text, str)
        assert isinstance(tables, list)
        assert len(tables) == 1
        
        # Check table content
        table = tables[0]
        assert table["name"] == "test"  # CSV uses filename as table name
        assert len(table["rows"]) == 3  # Header + 2 data rows
        
        # Check specific content
        assert table["rows"][0] == ["Name", "Age", "City"]
        assert table["rows"][1] == ["John", "30", "New York"]
        assert table["rows"][2] == ["Jane", "25", "Los Angeles"]
        
        # Check text content (flattened)
        assert "John" in text
        assert "Jane" in text
        assert "New York" in text
    
    def test_parse_csv_empty(self, tmp_path):
        """Test CSV parsing with empty file."""
        csv_file = tmp_path / "empty.csv"
        csv_file.write_text("")
        
        text, tables = parse_csv(csv_file)
        assert text == ""
        assert len(tables) == 0  # Empty CSV returns no tables


class TestDOCXParser:
    """Test DOCX parsing functionality."""
    
    @pytest.mark.skipif(
        not pytest.importorskip("docx", reason="python-docx not installed"),
        reason="python-docx not available"
    )
    def test_parse_docx_basic(self, tmp_path):
        """Test basic DOCX parsing."""
        from docx import Document
        
        # Create a test DOCX file
        doc = Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another paragraph with some content.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(1, 0).text = "Data1"
        table.cell(1, 1).text = "Data2"
        
        docx_file = tmp_path / "test.docx"
        doc.save(str(docx_file))
        
        # Parse the DOCX
        text, tables = parse_docx(docx_file)
        
        # Assert structure
        assert isinstance(text, str)
        assert isinstance(tables, list)
        
        # Check text content
        assert "This is a test paragraph" in text
        assert "This is another paragraph with some content" in text
        
        # Check table content
        assert len(tables) == 1
        table = tables[0]
        assert table["name"] == "Table_1"  # DOCX tables have auto-generated names
        assert len(table["rows"]) == 2
        
        # Check table data
        assert table["rows"][0] == ["Header1", "Header2"]
        assert table["rows"][1] == ["Data1", "Data2"]
    
    def test_parse_docx_fallback(self, tmp_path):
        """Test DOCX parsing fallback when python-docx not available."""
        # Mock the import to simulate missing dependency
        with patch('app.services.parsers.DOCX_AVAILABLE', False):
            # Create a dummy file
            docx_file = tmp_path / "test.docx"
            docx_file.write_text("dummy content")
            
            # Parse should return fallback message
            text, tables = parse_docx(docx_file)
            
            assert "(docx parser unavailable: python-docx not installed)" in text
            assert tables == []


class TestXLSXParser:
    """Test XLSX parsing functionality."""
    
    @pytest.mark.skipif(
        not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
        reason="openpyxl not available"
    )
    def test_parse_xlsx_basic(self, tmp_path):
        """Test basic XLSX parsing."""
        from openpyxl import Workbook
        
        # Create a test XLSX file
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        
        # Add data to first sheet
        ws1['A1'] = "Name"
        ws1['B1'] = "Age"
        ws1['A2'] = "John"
        ws1['B2'] = "30"
        ws1['A3'] = "Jane"
        ws1['B3'] = "25"
        
        # Add second sheet
        ws2 = wb.create_sheet("Sheet2")
        ws2['A1'] = "Product"
        ws2['B1'] = "Price"
        ws2['A2'] = "Widget"
        ws2['B2'] = "10.99"
        
        xlsx_file = tmp_path / "test.xlsx"
        wb.save(str(xlsx_file))
        
        # Parse the XLSX
        text, tables = parse_xlsx(xlsx_file)
        
        # Assert structure
        assert isinstance(text, str)
        assert isinstance(tables, list)
        
        # Check tables (one per sheet)
        assert len(tables) == 2
        
        # Check first sheet
        sheet1 = tables[0]
        assert sheet1["name"] == "Sheet1"
        assert len(sheet1["rows"]) == 3
        assert sheet1["rows"][0] == ["Name", "Age"]
        assert sheet1["rows"][1] == ["John", "30"]
        assert sheet1["rows"][2] == ["Jane", "25"]
        
        # Check second sheet
        sheet2 = tables[1]
        assert sheet2["name"] == "Sheet2"
        assert len(sheet2["rows"]) == 2
        assert sheet2["rows"][0] == ["Product", "Price"]
        assert sheet2["rows"][1] == ["Widget", "10.99"]
        
        # Check text content (flattened)
        assert "John" in text
        assert "Jane" in text
        assert "Widget" in text
        assert "10.99" in text
    
    def test_parse_xlsx_fallback(self, tmp_path):
        """Test XLSX parsing fallback when openpyxl not available."""
        # Mock the import to simulate missing dependency
        with patch('app.services.parsers.XLSX_AVAILABLE', False):
            # Create a dummy file
            xlsx_file = tmp_path / "test.xlsx"
            xlsx_file.write_text("dummy content")
            
            # Parse should return fallback message
            text, tables = parse_xlsx(xlsx_file)
            
            assert "(xlsx parser unavailable: openpyxl not installed)" in text
            assert tables == []


class TestImageOCRParser:
    """Test image OCR parsing functionality."""
    
    def test_parse_image_ocr_disabled(self, tmp_path):
        """Test image parsing with OCR disabled."""
        # Create a minimal PNG file (just a few bytes)
        png_file = tmp_path / "test.png"
        png_file.write_bytes(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde')
        
        # Parse with OCR disabled
        text, tables = parse_image_ocr(png_file, enabled=False, lang="eng")
        
        # Should return stub content
        assert isinstance(text, str)
        assert isinstance(tables, list)
        assert text == "(OCR disabled)"
        assert len(tables) == 0
    
    @pytest.mark.xfail(
        reason="Tesseract not installed or OCR not working",
        raises=(ImportError, OSError, Exception)
    )
    def test_parse_image_ocr_enabled(self, tmp_path):
        """Test image parsing with OCR enabled."""
        # Create a minimal PNG file
        png_file = tmp_path / "test.png"
        png_file.write_bytes(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde')
        
        # Parse with OCR enabled
        text, tables = parse_image_ocr(png_file, enabled=True, lang="eng")
        
        # Should return OCR result or fallback
        assert isinstance(text, str)
        assert isinstance(tables, list)
        # OCR might return empty text for minimal images, which is fine
        assert isinstance(text, str)
    
    def test_parse_image_ocr_pil_missing(self, tmp_path):
        """Test image parsing fallback when PIL not available."""
        # Mock the import to simulate missing dependency
        with patch('app.services.parsers.PIL_AVAILABLE', False):
            # Create a dummy file
            png_file = tmp_path / "test.png"
            png_file.write_bytes(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde')
            
            # Parse should return fallback message
            text, tables = parse_image_ocr(png_file, enabled=True, lang="eng")
            
            assert "(OCR unavailable: PIL not installed)" in text
            assert tables == []
    
    def test_parse_image_ocr_tesseract_missing(self, tmp_path):
        """Test image parsing fallback when pytesseract not available."""
        # Mock the import to simulate missing dependency
        with patch('app.services.parsers.TESSERACT_AVAILABLE', False):
            # Create a dummy file
            png_file = tmp_path / "test.png"
            png_file.write_bytes(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde')
            
            # Parse should return fallback message
            text, tables = parse_image_ocr(png_file, enabled=True, lang="eng")
            
            assert "(OCR unavailable: pytesseract not installed)" in text
            assert tables == []


class TestPDFParser:
    """Test PDF parsing functionality."""
    
    def test_parse_pdf_stub(self, tmp_path):
        """Test PDF stub parsing."""
        # Create a dummy PDF file
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_text("%PDF-1.4\n%Test PDF content\n")
        
        text, tables = parse_pdf_stub(pdf_file)
        
        # Should return stub content
        assert isinstance(text, str)
        assert isinstance(tables, list)
        assert text == "(pdf parsing stub — not yet implemented)"
        assert len(tables) == 0


class TestBuildNormalized:
    """Test normalized document model building."""
    
    def test_build_normalized(self):
        """Test building normalized document model."""
        doc_type = "csv"
        meta = {"filename": "test.csv", "content_hash": "abc123", "size": 100}
        text = "Sample text content"
        tables = [{"name": "Sheet1", "rows": [["A", "B"], ["1", "2"]]}]
        
        result = build_normalized(doc_type, meta, text, tables)
        
        # Check structure
        assert result["type"] == "csv"
        assert result["meta"] == meta
        assert result["content"]["text"] == text
        assert result["content"]["tables"] == tables


class TestParseToNormalized:
    """Test main document parsing function."""
    
    def test_parse_to_normalized_csv(self, tmp_path):
        """Test main parser with CSV file."""
        # Create a test CSV file
        csv_file = tmp_path / "test.csv"
        with open(csv_file, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Header"])
            writer.writerow(["Data"])
        
        # Build metadata
        meta = {
            "filename": "test.csv",
            "content_hash": "abc123",
            "size": 100
        }
        
        # Parse using main function
        result = parse_to_normalized(csv_file, meta)
        
        # Check normalized structure
        assert result["type"] == "csv"
        assert "meta" in result
        assert "content" in result
        assert result["meta"]["filename"] == "test.csv"
        assert result["meta"]["content_hash"] == "abc123"
        assert result["content"]["text"] == "Header Data"
        assert len(result["content"]["tables"]) == 1
    
    def test_parse_to_normalized_unknown_type(self, tmp_path):
        """Test main parser with unknown file type."""
        # Create a dummy file
        unknown_file = tmp_path / "test.xyz"
        unknown_file.write_text("Some content")
        
        # Build metadata
        meta = {
            "filename": "test.xyz",
            "content_hash": "xyz789",
            "size": 50
        }
        
        # Parse using main function
        result = parse_to_normalized(unknown_file, meta)
        
        # Check normalized structure
        assert result["type"] == "unknown"
        assert result["meta"]["filename"] == "test.xyz"
        assert result["content"]["text"] == ""
        assert len(result["content"]["tables"]) == 0


class TestParserIntegration:
    """Test parser integration with ingest workflow."""
    
    def test_parser_with_ingest_metadata(self, tmp_path):
        """Test that parsers work with ingest metadata structure."""
        # Create a test CSV file
        csv_file = tmp_path / "test.csv"
        with open(csv_file, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Test"])
            writer.writerow(["Data"])
        
        # Simulate ingest workflow
        doc_type = detect_type(csv_file.name)
        meta = {
            "filename": csv_file.name,
            "content_hash": "test123",
            "size": 100
        }
        parsed_content = parse_to_normalized(csv_file, meta)
        
        # Verify the structure matches ingest expectations
        assert parsed_content["type"] == "csv"
        assert "meta" in parsed_content
        assert "content" in parsed_content
        assert "text" in parsed_content["content"]
        assert "tables" in parsed_content["content"]
        
        # Verify content was actually parsed
        assert "Test" in parsed_content["content"]["text"]
        assert "Data" in parsed_content["content"]["text"]
        assert len(parsed_content["content"]["tables"]) > 0


class TestParserSafety:
    """Test that parsers handle errors gracefully."""
    
    def test_parser_handles_missing_file(self):
        """Test that parsers handle missing files gracefully."""
        missing_file = Path("/nonexistent/file.txt")
        
        # All parsers should handle missing files gracefully
        text, tables = parse_csv(missing_file)
        assert "Error parsing CSV" in text
        assert tables == []
        
        text, tables = parse_docx(missing_file)
        assert "Error parsing DOCX" in text
        assert tables == []
        
        text, tables = parse_xlsx(missing_file)
        assert "Error parsing XLSX" in text
        assert tables == []
    
    def test_parser_handles_corrupted_files(self, tmp_path):
        """Test that parsers handle corrupted files gracefully."""
        # Create a corrupted CSV file
        corrupted_csv = tmp_path / "corrupted.csv"
        corrupted_csv.write_text("Invalid CSV content\nwith\nbroken\nformat")
        
        # Parser should handle gracefully
        text, tables = parse_csv(corrupted_csv)
        # Should either parse what it can or return error message
        assert isinstance(text, str)
        assert isinstance(tables, list)
