#!/usr/bin/env python3
"""
Demo seeder script for EstimAI.
Generates sample files for demonstration purposes.
"""

import json
import os
from pathlib import Path
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

# Sample file definitions
SAMPLE_FILES = [
    {
        "slug": "ti-bid-pdf",
        "filename": "small_ti_bid_package.pdf",
        "name": "Small TI Bid Package (PDF, 6 pages)",
        "mime": "application/pdf",
        "generator": "generate_pdf"
    },
    {
        "slug": "unit-takeoff-csv",
        "filename": "unit_takeoff_200u.csv",
        "name": "Unit Takeoff CSV (Multifamily 200u)",
        "mime": "text/csv",
        "generator": "generate_csv"
    },
    {
        "slug": "spec-excerpt-docx",
        "filename": "spec_book_excerpt.docx",
        "name": "Spec Book Excerpt (DOCX)",
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "generator": "generate_docx"
    },
    {
        "slug": "estimate-worksheet-xlsx",
        "filename": "estimate_worksheet.xlsx",
        "name": "Estimate Worksheet (XLSX)",
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "generator": "generate_xlsx"
    },
    {
        "slug": "site-plan-png",
        "filename": "site_plan.png",
        "name": "Site Plan (PNG)",
        "mime": "image/png",
        "generator": "generate_png"
    }
]


def ensure_samples_directory() -> Path:
    """Ensure the samples directory exists."""
    app_dir = Path(__file__).resolve().parent.parent
    samples_dir = app_dir / "samples"
    samples_dir.mkdir(exist_ok=True)
    return samples_dir


def generate_pdf(file_path: Path) -> int:
    """Generate a small TI bid package PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(str(file_path), pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12
        )
        
        story = []
        
        # Page 1: Title Page
        story.append(Paragraph("TENANT IMPROVEMENT BID PACKAGE", title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("Project Overview", heading_style))
        story.append(Paragraph("This document outlines the scope of work for tenant improvements including framing, drywall, flooring, and fixtures. The project involves approximately 200 units with standard finishes and modern amenities.", styles['Normal']))
        story.append(PageBreak())
        
        # Page 2: Scope of Work
        story.append(Paragraph("Scope of Work", title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Framing & Structure", heading_style))
        story.append(Paragraph("• Interior wall framing with 2x4 studs at 16\" on center", styles['Normal']))
        story.append(Paragraph("• Door and window rough openings", styles['Normal']))
        story.append(Paragraph("• Ceiling grid support structure", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Drywall & Finishes", heading_style))
        story.append(Paragraph("• 5/8\" drywall on walls and ceilings", styles['Normal']))
        story.append(Paragraph("• Joint compound and tape finish", styles['Normal']))
        story.append(Paragraph("• Primer coat application", styles['Normal']))
        story.append(PageBreak())
        
        # Page 3: Materials
        story.append(Paragraph("Materials & Specifications", title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Flooring", heading_style))
        story.append(Paragraph("• Luxury vinyl plank flooring in living areas", styles['Normal']))
        story.append(Paragraph("• Ceramic tile in bathrooms and kitchens", styles['Normal']))
        story.append(Paragraph("• Carpet in bedrooms", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Fixtures", heading_style))
        story.append(Paragraph("• Standard plumbing fixtures", styles['Normal']))
        story.append(Paragraph("• Electrical outlets and switches", styles['Normal']))
        story.append(Paragraph("• HVAC registers and returns", styles['Normal']))
        story.append(PageBreak())
        
        # Page 4: Timeline
        story.append(Paragraph("Project Timeline", title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Phase 1: Demolition & Prep (Week 1-2)", heading_style))
        story.append(Paragraph("• Remove existing finishes and fixtures", styles['Normal']))
        story.append(Paragraph("• Structural modifications if required", styles['Normal']))
        story.append(Paragraph("• Rough-in preparation", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Phase 2: Construction (Week 3-8)", heading_style))
        story.append(Paragraph("• Framing and structural work", styles['Normal']))
        story.append(Paragraph("• Mechanical, electrical, and plumbing", styles['Normal']))
        story.append(Paragraph("• Drywall and finishing", styles['Normal']))
        story.append(PageBreak())
        
        # Page 5: Cost Breakdown
        story.append(Paragraph("Cost Breakdown", title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Labor Costs", heading_style))
        story.append(Paragraph("• Framing: $45,000", styles['Normal']))
        story.append(Paragraph("• Drywall: $38,000", styles['Normal']))
        story.append(Paragraph("• Flooring: $52,000", styles['Normal']))
        story.append(Paragraph("• Fixtures: $28,000", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Material Costs", heading_style))
        story.append(Paragraph("• Lumber and fasteners: $18,000", styles['Normal']))
        story.append(Paragraph("• Drywall and finishing: $12,000", styles['Normal']))
        story.append(Paragraph("• Flooring materials: $35,000", styles['Normal']))
        story.append(PageBreak())
        
        # Page 6: Summary
        story.append(Paragraph("Project Summary", title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Total Project Cost: $266,000", heading_style))
        story.append(Paragraph("Project Duration: 8 weeks", styles['Normal']))
        story.append(Paragraph("Units Affected: 200", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("This bid package provides a comprehensive overview of the tenant improvement project. All work will be completed according to local building codes and industry standards.", styles['Normal']))
        
        doc.build(story)
        
        # Return file size
        return file_path.stat().st_size
        
    except ImportError:
        logger.warning("reportlab not available, creating stub PDF")
        with open(file_path, 'w') as f:
            f.write("PDF stub - install reportlab for full PDF generation")
        return file_path.stat().st_size


def generate_csv(file_path: Path) -> int:
    """Generate a unit takeoff CSV file."""
    csv_content = """item,quantity,unit,note
2x4 Studs,2400,ea,16" on center
5/8" Drywall,120,sheet,4x8 sheets
Joint Compound,60,gal,20lb buckets
Drywall Tape,120,roll,500ft rolls
Luxury Vinyl Plank,8000,sqft,6" x 36" planks
Ceramic Tile,1200,sqft,12" x 12" tiles
Carpet,2400,sqft,berber style
Paint Primer,40,gal,5 gallon buckets
Paint Finish,40,gal,eggshell finish
Plumbing Fixtures,200,ea,standard fixtures
Electrical Outlets,400,ea,duplex receptacles
Light Fixtures,200,ea,LED flush mount
HVAC Registers,200,ea,6" x 12" registers
Door Hardware,200,ea,passage sets
Window Treatments,200,ea,mini blinds
Baseboard,800,lf,3-1/4" height
Crown Molding,800,lf,2-1/2" profile
Cabinet Hardware,200,ea,knobs and pulls
Mirrors,200,ea,24" x 36" bathroom mirrors
Smoke Detectors,200,ea,hardwired with battery backup"""
    
    with open(file_path, 'w') as f:
        f.write(csv_content)
    
    return file_path.stat().st_size


def generate_docx(file_path: Path) -> int:
    """Generate a spec book excerpt DOCX file."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Specification Book Excerpt', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        doc.add_heading('Project Specifications', level=1)
        doc.add_paragraph('This specification book excerpt outlines the requirements for the multifamily tenant improvement project. All materials and workmanship must meet or exceed the standards specified herein.')
        
        # Section 1: General Requirements
        doc.add_heading('1.0 General Requirements', level=1)
        doc.add_paragraph('1.1 Scope of Work')
        doc.add_paragraph('The contractor shall provide all labor, materials, equipment, and services necessary to complete the tenant improvement work as described in these specifications.')
        
        doc.add_paragraph('1.2 Quality Assurance')
        doc.add_paragraph('All work shall be performed in accordance with local building codes, manufacturer specifications, and industry best practices.')
        
        # Section 2: Materials
        doc.add_heading('2.0 Materials', level=1)
        doc.add_paragraph('2.1 Framing Materials')
        doc.add_paragraph('• Studs: 2x4 kiln-dried spruce-pine-fir, minimum grade #2')
        doc.add_paragraph('• Plates: 2x4 kiln-dried spruce-pine-fir, minimum grade #2')
        doc.add_paragraph('• Fasteners: 3-1/2" common nails or 3" drywall screws')
        
        doc.add_paragraph('2.2 Drywall Materials')
        doc.add_paragraph('• Type X fire-rated drywall, 5/8" thickness')
        doc.add_paragraph('• Joint compound: pre-mixed, lightweight, ready-mixed')
        doc.add_paragraph('• Joint tape: paper tape, 2" width')
        
        # Section 3: Installation
        doc.add_heading('3.0 Installation', level=1)
        doc.add_paragraph('3.1 Framing Installation')
        doc.add_paragraph('• Studs shall be installed 16" on center, maximum')
        doc.add_paragraph('• Plates shall be securely fastened to structural elements')
        doc.add_paragraph('• Door and window openings shall be properly reinforced')
        
        doc.add_paragraph('3.2 Drywall Installation')
        doc.add_paragraph('• Sheets shall be installed with staggered joints')
        doc.add_paragraph('• All joints shall be taped and finished to level 4')
        doc.add_paragraph('• Screws shall be countersunk and not over-driven')
        
        # Section 4: Finishing
        doc.add_heading('4.0 Finishing', level=1)
        doc.add_paragraph('4.1 Paint Preparation')
        doc.add_paragraph('• All surfaces shall be clean and free of dust')
        doc.add_paragraph('• Primer shall be applied before finish coats')
        doc.add_paragraph('• Minimum of two finish coats required')
        
        doc.add_paragraph('4.2 Final Inspection')
        doc.add_paragraph('• All work shall be inspected by the project manager')
        doc.add_paragraph('• Punch list items shall be completed before final acceptance')
        doc.add_paragraph('• As-built drawings shall be provided upon completion')
        
        doc.save(str(file_path))
        return file_path.stat().st_size
        
    except ImportError:
        logger.warning("python-docx not available, creating stub DOCX")
        with open(file_path, 'w') as f:
            f.write("DOCX stub - install python-docx for full DOCX generation")
        return file_path.stat().st_size


def generate_xlsx(file_path: Path) -> int:
    """Generate an estimate worksheet XLSX file."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
        from openpyxl.utils import get_column_letter
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Estimate Worksheet"
        
        # Headers
        headers = ["Cost Code", "Description", "Qty", "Unit", "Unit Cost", "Total"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
        
        # Data rows
        data = [
            ["01-100", "Mobilization", 1, "LS", 2500, "=D2*E2"],
            ["02-100", "Demolition", 8000, "SF", 3.50, "=D3*E3"],
            ["03-100", "Framing", 8000, "SF", 8.75, "=D4*E4"],
            ["04-100", "Drywall", 8000, "SF", 6.25, "=D5*E5"],
            ["05-100", "Flooring", 8000, "SF", 12.50, "=D6*E6"],
            ["06-100", "Paint", 8000, "SF", 3.25, "=D7*E7"],
            ["07-100", "Plumbing", 200, "EA", 150.00, "=D8*E8"],
            ["08-100", "Electrical", 200, "EA", 200.00, "=D9*E9"],
            ["09-100", "HVAC", 200, "EA", 300.00, "=D10*E10"],
            ["10-100", "Finishes", 200, "EA", 500.00, "=D11*E11"]
        ]
        
        for row, row_data in enumerate(data, 2):
            for col, value in enumerate(row_data, 1):
                ws.cell(row=row, column=col, value=value)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(str(file_path))
        return file_path.stat().st_size
        
    except ImportError:
        logger.warning("openpyxl not available, creating stub XLSX")
        with open(file_path, 'w') as f:
            f.write("XLSX stub - install openpyxl for full XLSX generation")
        return file_path.stat().st_size


def generate_png(file_path: Path) -> int:
    """Generate a site plan PNG image."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create 800x600 white image
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        # Draw black border
        draw.rectangle([(0, 0), (799, 599)], outline='black', width=3)
        
        # Draw gray rectangle for "Building A"
        building_rect = [(100, 150), (400, 450)]
        draw.rectangle(building_rect, fill='lightgray', outline='black', width=2)
        
        # Add "Building A" label
        try:
            # Try to use a default font
            font = ImageFont.load_default()
        except:
            font = None
        
        # Center the text in the building rectangle
        text = "Building A"
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        text_x = building_rect[0][0] + (building_rect[1][0] - building_rect[0][0] - text_width) // 2
        text_y = building_rect[0][1] + (building_rect[1][1] - building_rect[0][1] - text_height) // 2
        
        draw.text((text_x, text_y), text, fill='black', font=font)
        
        # Draw some additional site elements
        # Parking lot
        draw.rectangle([(450, 200), (700, 350)], fill='lightblue', outline='black', width=1)
        draw.text((520, 270), "Parking", fill='black', font=font)
        
        # Landscaping
        draw.ellipse([(50, 50), (150, 100)], fill='lightgreen', outline='darkgreen', width=1)
        draw.text((80, 70), "Trees", fill='darkgreen', font=font)
        
        # Legend
        legend_y = 500
        draw.text((50, legend_y), "Legend:", fill='black', font=font)
        draw.rectangle([(50, legend_y + 20), (70, legend_y + 40)], fill='lightgray', outline='black')
        draw.text((80, legend_y + 20), "Buildings", fill='black', font=font)
        draw.rectangle([(50, legend_y + 50), (70, legend_y + 70)], fill='lightblue', outline='black')
        draw.text((80, legend_y + 50), "Parking", fill='black', font=font)
        draw.ellipse([(50, legend_y + 80), (70, legend_y + 100)], fill='lightgreen', outline='darkgreen')
        draw.text((80, legend_y + 80), "Landscaping", fill='darkgreen', font=font)
        
        img.save(file_path, 'PNG')
        return file_path.stat().st_size
        
    except ImportError:
        logger.warning("Pillow not available, creating stub PNG")
        with open(file_path, 'w') as f:
            f.write("PNG stub - install Pillow for full PNG generation")
        return file_path.stat().st_size


def generate_sample_file(sample_info: Dict[str, Any], samples_dir: Path) -> Dict[str, Any]:
    """Generate a single sample file and return its metadata."""
    file_path = samples_dir / sample_info["filename"]
    
    # Skip if file already exists
    if file_path.exists():
        logger.info(f"Sample file already exists: {sample_info['filename']}")
        return {
            "slug": sample_info["slug"],
            "filename": sample_info["filename"],
            "name": sample_info["name"],
            "bytes": file_path.stat().st_size,
            "mime": sample_info["mime"]
        }
    
    # Generate the file
    generator_name = sample_info["generator"]
    generator_func = globals()[generator_name]
    
    logger.info(f"Generating sample file: {sample_info['filename']}")
    file_size = generator_func(file_path)
    
    return {
        "slug": sample_info["slug"],
        "filename": sample_info["filename"],
        "name": sample_info["name"],
        "bytes": file_size,
        "mime": sample_info["mime"]
    }


def run() -> None:
    """Main function to run the demo seeder."""
    logger.info("🌱 Starting demo sample file generation...")
    
    # Ensure samples directory exists
    samples_dir = ensure_samples_directory()
    logger.info(f"📁 Samples directory: {samples_dir}")
    
    # Generate all sample files
    generated_files = []
    for sample_info in SAMPLE_FILES:
        try:
            file_metadata = generate_sample_file(sample_info, samples_dir)
            generated_files.append(file_metadata)
            logger.info(f"✅ Generated: {sample_info['filename']} ({file_metadata['bytes']} bytes)")
        except Exception as e:
            logger.error(f"❌ Failed to generate {sample_info['filename']}: {e}")
            # Add stub entry for failed generation
            generated_files.append({
                "slug": sample_info["slug"],
                "filename": sample_info["filename"],
                "name": sample_info["name"],
                "bytes": 0,
                "mime": sample_info["mime"]
            })
    
    # Write index.json
    index_path = samples_dir / "index.json"
    with open(index_path, 'w') as f:
        json.dump(generated_files, f, indent=2)
    
    logger.info(f"📋 Sample index written to: {index_path}")
    logger.info(f"🎉 Demo seeding completed! Generated {len(generated_files)} sample files")


if __name__ == "__main__":
    # Set up basic logging
    logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
    run()
