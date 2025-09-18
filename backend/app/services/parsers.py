"""
Dependency-safe document parsers with graceful degradation.

This module provides parsers for various document formats that gracefully
handle missing dependencies by returning informative stub content.
"""

import csv
import logging
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional

logger = logging.getLogger(__name__)

# Safe import checks with informative fallbacks
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX parsing will be stubbed")

try:
    from openpyxl import load_workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logger.warning("openpyxl not available - XLSX parsing will be stubbed")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not available - Image processing will be stubbed")

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract not available - OCR will be stubbed")


def _safe_str(value: Any) -> str:
    """Safely convert any value to string, handling None and empty values."""
    if value is None:
        return ""
    return str(value).strip()


def _mk_table(name: Optional[str], rows: List[List[str]]) -> Dict[str, Any]:
    """Create a table structure with safe string handling."""
    return {
        "name": name,
        "rows": [[_safe_str(cell) for cell in row] for row in rows]
    }


def detect_type(filename: str) -> str:
    """
    Detect document type based on filename extension.
    
    Args:
        filename: Name of the file
        
    Returns:
        Document type string: docx, xlsx, csv, image, pdf, or unknown
    """
    if not filename:
        return "unknown"
    
    ext = Path(filename).suffix.lower()
    
    if ext == '.docx':
        return "docx"
    elif ext == '.xlsx':
        return "xlsx"
    elif ext == '.csv':
        return "csv"
    elif ext in ['.png', '.jpg', '.jpeg', '.tif', '.tiff']:
        return "image"
    elif ext == '.pdf':
        return "pdf"
    else:
        return "unknown"


def parse_docx(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Parse DOCX file and extract text and tables.
    
    Args:
        path: Path to the DOCX file
        
    Returns:
        Tuple of (text, tables) where tables is list of {name, rows}
    """
    if not DOCX_AVAILABLE:
        return (
            "(docx parser unavailable: python-docx not installed)",
            []
        )
    
    try:
        doc = Document(path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            tables.append(_mk_table(f"Table_{i+1}", table_data))
        
        return "\n".join(text_parts), tables
        
    except Exception as e:
        logger.error(f"Failed to parse DOCX file {path}: {e}")
        return f"(Error parsing DOCX: {e})", []


def parse_xlsx(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Parse XLSX file and extract text and tables from each sheet.
    
    Args:
        path: Path to the XLSX file
        
    Returns:
        Tuple of (text, tables) where tables is list of {name, rows}
    """
    if not XLSX_AVAILABLE:
        return (
            "(xlsx parser unavailable: openpyxl not installed)",
            []
        )
    
    try:
        workbook = load_workbook(path, data_only=True)
        
        # Extract text from all cells (flattened)
        text_parts = []
        tables = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Extract table data from sheet
            table_data = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None and str(cell).strip() for cell in row):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    table_data.append(row_data)
            
            if table_data:
                tables.append(_mk_table(sheet_name, table_data))
                
                # Add sheet content to text (flattened)
                for row in table_data:
                    text_parts.extend([str(cell) for cell in row if str(cell).strip()])
        
        return " ".join(text_parts), tables
        
    except Exception as e:
        logger.error(f"Failed to parse XLSX file {path}: {e}")
        return f"(Error parsing XLSX: {e})", []


def parse_csv(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Parse CSV file and extract text and table data.
    
    Args:
        path: Path to the CSV file
        
    Returns:
        Tuple of (text, tables) where tables is list of {name, rows}
    """
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as file:
            reader = csv.reader(file)
            rows = list(reader)
        
        # Extract text (flattened)
        text_parts = []
        for row in rows:
            text_parts.extend([str(cell) for cell in row if str(cell).strip()])
        
        # Create table structure
        tables = []
        if rows:
            tables.append(_mk_table(Path(path).stem, rows))
        
        return " ".join(text_parts), tables
        
    except Exception as e:
        logger.error(f"Failed to parse CSV file {path}: {e}")
        return f"(Error parsing CSV: {e})", []


def parse_image_ocr(path: Path, enabled: bool, lang: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Parse image file with OCR if enabled, otherwise return stub.
    
    Args:
        path: Path to the image file
        enabled: Whether OCR is enabled
        lang: OCR language code
        
    Returns:
        Tuple of (text, tables) where tables is always empty for images
    """
    if not enabled:
        return "(OCR disabled)", []
    
    if not PIL_AVAILABLE:
        return "(OCR unavailable: PIL not installed)", []
    
    if not TESSERACT_AVAILABLE:
        return "(OCR unavailable: pytesseract not installed)", []
    
    try:
        # Open image with PIL
        image = Image.open(path)
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang=lang)
        
        # Clean up text
        text = text.strip()
        
        return text if text else "(No text detected)", []
        
    except Exception as e:
        logger.error(f"Failed to parse image with OCR {path}: {e}")
        return f"(Error processing image: {e})", []


def parse_pdf_stub(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Stub PDF parser (placeholder for future PDF text extraction).
    
    Args:
        path: Path to the PDF file
        
    Returns:
        Tuple of stub content and empty tables
    """
    return "(pdf parsing stub — not yet implemented)", []


def parse_unknown(path: Path) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Parser for unknown file types.
    
    Args:
        path: Path to the unknown file
        
    Returns:
        Tuple of empty text and empty tables
    """
    return "", []


def build_normalized(doc_type: str, meta: Dict[str, Any], text: str, tables: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build normalized document model.
    
    Args:
        doc_type: Document type identifier
        meta: File metadata (filename, content_hash, size)
        text: Extracted text content
        tables: Extracted table data
        
    Returns:
        Normalized document model
    """
    return {
        "type": doc_type,
        "meta": meta,
        "content": {
            "text": text,
            "tables": tables
        }
    }


def parse_to_normalized(path: Path, meta: Dict[str, Any], ocr_enabled: bool = False, ocr_lang: str = "eng") -> Dict[str, Any]:
    """
    Parse document to normalized format with auto-detection.
    
    Args:
        path: Path to the document
        meta: File metadata (filename, content_hash, size)
        ocr_enabled: Whether OCR is enabled for images
        ocr_lang: OCR language code
        
    Returns:
        Normalized document model: { type, meta, content: { text, tables } }
    """
    # Auto-detect document type
    doc_type = detect_type(path.name)
    
    # Parse based on type
    if doc_type == "docx":
        text, tables = parse_docx(path)
    elif doc_type == "xlsx":
        text, tables = parse_xlsx(path)
    elif doc_type == "csv":
        text, tables = parse_csv(path)
    elif doc_type == "image":
        text, tables = parse_image_ocr(path, ocr_enabled, ocr_lang)
    elif doc_type == "pdf":
        text, tables = parse_pdf_stub(path)
    else:
        text, tables = parse_unknown(path)
    
    # Return normalized model
    return build_normalized(doc_type, meta, text, tables)


# Legacy function names for backward compatibility
def parse_document(path: Path, doc_type: str, ocr_enabled: bool = False, ocr_lang: str = "eng") -> Dict[str, Any]:
    """
    Legacy function for backward compatibility.
    
    Args:
        path: Path to the document
        doc_type: Document type
        ocr_enabled: Whether OCR is enabled for images
        ocr_lang: OCR language code
        
    Returns:
        Normalized document model
    """
    # Build metadata from path
    meta = {
        "filename": path.name,
        "content_hash": "",  # Will be set by caller
        "size": path.stat().st_size if path.exists() else 0
    }
    
    return parse_to_normalized(path, meta, ocr_enabled, ocr_lang)
